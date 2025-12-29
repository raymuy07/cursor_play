"""
CV Embedding Script
Extracts text from CV files (PDF or DOCX), transforms via LLM, and generates embeddings.
"""

import os
import sys
import re
import logging
from pathlib import Path
from typing import Dict
import yaml

# Add parent directory to path for imports
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from common.utils import load_config, TextEmbedder
from openai import OpenAI

logger = logging.getLogger(__name__)

# Third-party imports for document processing
try:
    import PyPDF2
except ImportError:
    PyPDF2 = None

try:
    from docx import Document
except ImportError:
    Document = None

# Default prompts path
PROMPTS_PATH = Path(__file__).parent.parent / "prompts" / "job_matching_v1.yaml"


class CVProcessor:
    """
    Full CV pipeline: extract → clean → transform (LLM) → embed.
    """

    def __init__(self):
        self.embedder = TextEmbedder()
        self.config = load_config()
        self.client = OpenAI(api_key=self.config.get('openai_api_key'))
        self.prompts = self._load_prompts()

    def process(self, file_path: str) -> dict:
        """
        Full pipeline: extract → clean → transform via LLM → embed.
        Returns dict with embedding and metadata.
        """
        # 1. Extract raw text from CV
        raw_text = self._extract_text(file_path)
        raw_text = self._clean_text(raw_text)
        logger.info(f"Extracted {len(raw_text)} chars from {Path(file_path).name}")

        # 2. Transform CV to JD-style description via LLM
        transformed_text = self._transform_with_llm(raw_text)
        logger.info(f"Transformed to {len(transformed_text)} chars via LLM")

        # 3. Embed the transformed text
        result = self.embedder.embed(transformed_text)
        result['source_file'] = Path(file_path).name
        result['transformed_text'] = transformed_text

        return result

    def _load_prompts(self) -> dict:
        """Load prompts from YAML file."""
        with open(PROMPTS_PATH, 'r', encoding='utf-8') as f:
            return yaml.safe_load(f)

    def _transform_with_llm(self, cv_text: str) -> str:
        """Transform CV text to JD-style description using LLM."""
        prompt_config = self.prompts['cv_to_job_description']

        system_prompt = prompt_config['system_prompt']
        user_message = prompt_config['user_template'].format(cv_text=cv_text)

        response = self.client.chat.completions.create(
            model=prompt_config.get('model', 'gpt-4'),
            temperature=prompt_config.get('temperature', 0.3),
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_message}
            ]
        )

        return response.choices[0].message.content

    def _extract_text(self, file_path: str) -> str:
        """Auto-detect format and extract."""
        ext = Path(file_path).suffix.lower()

        if ext == '.pdf':
            return self._extract_text_from_pdf(file_path)
        elif ext == '.docx':
            return self._extract_text_from_docx(file_path)
        else:
            raise ValueError(f"Unsupported format: {ext}. Use .pdf or .docx")

    def _extract_text_from_pdf(self, file_path: str) -> str:

        if PyPDF2 is None:
            raise ImportError("PyPDF2 is required for PDF processing. Install it with: pip install PyPDF2")

        text_parts = []

        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                num_pages = len(pdf_reader.pages)

                logger.info(f"Reading PDF with {num_pages} pages...")

                for page_num in range(num_pages):
                    page = pdf_reader.pages[page_num]
                    text = page.extract_text()
                    if text:
                        text_parts.append(text)

            full_text = "\n".join(text_parts)
            return self._clean_text(full_text)

        except Exception as e:
            raise RuntimeError(f"Error reading PDF file: {str(e)}")

    def _extract_text_from_docx(self, file_path: str) -> str:

        if Document is None:
            raise ImportError("python-docx is required for DOCX processing. Install it with: pip install python-docx")

        try:
            doc = Document(file_path)

            # Extract text from all paragraphs
            paragraphs = [paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()]

            # Also extract text from tables if any
            table_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        table_text.append(" | ".join(row_text))

            # Combine paragraphs and table text
            all_text = paragraphs + table_text
            full_text = "\n".join(all_text)

            logger.info(f"Read DOCX with {len(paragraphs)} paragraphs and {len(doc.tables)} tables")

            return self._clean_text(full_text)

        except Exception as e:
            raise RuntimeError(f"Error reading DOCX file: {str(e)}")

    def _clean_text(self, text: str) -> str:
        """Remove excessive whitespace and normalize text."""
        text = re.sub(r'\s+', ' ', text)
        text = re.sub(r'\n\s*\n', '\n\n', text)
        text = text.strip()
        return text


if __name__ == "__main__":
    from common.utils import setup_logging
    setup_logging()

    # Example usage
    processor = CVProcessor()
    result = processor.process("Guy Leiba-CV-0925.pdf")
    print(f"Embedding dim: {len(result['embedding'])}")
    print(f"Transformed text preview: {result['transformed_text'][:500]}...")


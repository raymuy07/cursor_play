
"""
CV manager service
Extracts text from CV files (PDF or DOCX), transforms via LLM, and generates embeddings.
"""

import logging
import re
from pathlib import Path

##can we use something internal?
import PyPDF2
from docx import Document

from app.common.txt_embedder import TextEmbedder
from app.common.txt_converter_ai import TextConverterAI

logger = logging.getLogger(__name__)


class CVManager:

    def __init__(self,text_embedder: TextEmbedder, text_converter_ai: TextConverterAI):
        self.text_embedder = text_embedder
        self.text_converter_ai = text_converter_ai

    def process(self, file_path: str) -> dict:
        """
        Full pipeline: extract → clean → transform via LLM → embed.
        Returns dict with embedding and metadata.
        """
        # 1. Extract raw text from CV
        raw_text = self._extract_text(file_path)
        raw_text = self._clean_text(raw_text)
        logger.info(f"Extracted {len(raw_text)} chars from {Path(file_path).name}")

        # 2. Transform CV to JD-style description via LLM
        transformed_text = self.text_converter_ai.run_through_llm(raw_text, "cv_to_job_description")
        logger.info(f"Transformed to {len(transformed_text)} chars via LLM")

        # 3. Embed the transformed text
        result = self.text_embedder.embed_immediate(transformed_text)
        result["source_file"] = Path(file_path).name
        result["transformed_text"] = transformed_text

        return result


    def _extract_text(self, file_path: str) -> str:
        """Auto-detect format and extract."""
        ext = Path(file_path).suffix.lower()

        if ext == ".pdf":
            return self._extract_text_from_pdf(file_path)
        elif ext == ".docx":
            return self._extract_text_from_docx(file_path)
        else:
            raise ValueError(f"Unsupported format: {ext}. Use .pdf or .docx")

    def _extract_text_from_pdf(self, file_path: str) -> str:
        if PyPDF2 is None:
            raise ImportError("PyPDF2 is required for PDF processing. Install it with: pip install PyPDF2")

        text_parts = []

        try:
            with open(file_path, "rb") as file:
                pdf_reader = PyPDF2.PdfReader(file)
                num_pages = len(pdf_reader.pages)

                logger.info(f"Reading PDF with {num_pages} pages...")

                for page_num in range(num_pages):
                    page = pdf_reader.pages[page_num]
                    text = page.extract_text()
                    if text:
                        text_parts.append(text)

            full_text = "\n".join(text_parts)
            return self._clean_text(full_text)

        except Exception as e:
            raise RuntimeError(f"Error reading PDF file: {e!s}")

    def _extract_text_from_docx(self, file_path: str) -> str:
        if Document is None:
            raise ImportError("python-docx is required for DOCX processing. Install it with: pip install python-docx")

        try:
            doc = Document(file_path)

            # Extract text from all paragraphs
            paragraphs = [paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()]

            # Also extract text from tables if any
            table_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        table_text.append(" | ".join(row_text))

            # Combine paragraphs and table text
            all_text = paragraphs + table_text
            full_text = "\n".join(all_text)

            logger.info(f"Read DOCX with {len(paragraphs)} paragraphs and {len(doc.tables)} tables")

            return self._clean_text(full_text)

        except Exception as e:
            raise RuntimeError(f"Error reading DOCX file: {e!s}")

    def _clean_text(self, text: str) -> str:
        """Remove excessive whitespace and normalize text."""
        text = re.sub(r"\s+", " ", text)
        text = re.sub(r"\n\s*\n", "\n\n", text)
        text = text.strip()
        return text
